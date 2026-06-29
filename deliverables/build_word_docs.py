from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "deliverables"

SOURCES = [
    ("01-产品源代码及可执行文件说明.md", "01-产品源代码及可执行文件说明.docx", "产品源代码及可执行文件说明"),
    ("02-产品部署和使用手册.md", "02-产品部署和使用手册.docx", "产品部署和使用手册"),
    ("03-产品总体设计文档.md", "03-产品总体设计文档.docx", "产品总体设计文档"),
]

FONT = "Microsoft YaHei"
MONO = "Consolas"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(96, 96, 96)
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F6F8FA"
BORDER = "D9E2EC"


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_footer(section, title):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_run_font(r, size=8.5, color=GRAY)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc, title, source_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("智游景行")
    set_run_font(r, size=12, bold=True, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, size=24, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("景区导览 AI 数字人系统交付文档")
    set_run_font(r, size=12, color=GRAY)

    rows = [
        ("文档类型", title),
        ("项目目录", str(ROOT)),
        ("源文件", source_name),
        ("生成格式", "Microsoft Word (.docx)"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1900, 7460])
    for i, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), LIGHT_GRAY)
        for j, text in enumerate((label, value)):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=(j == 0))

    doc.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in row):
            rows.append(row)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    base_widths = {
        2: [2700, 6660],
        3: [2000, 3200, 4160],
        4: [1700, 2300, 3000, 2360],
    }
    widths = base_widths.get(cols, [int(9360 / cols)] * cols)
    if sum(widths) != 9360:
        widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text.replace("`", ""))
            set_run_font(run, size=9.2 if cols >= 4 else 9.8, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph()


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for idx, line in enumerate(code.rstrip("\n").splitlines() or [""]):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name=MONO, size=9, color=RGBColor(40, 40, 40))
    doc.add_paragraph()


def add_rich_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_run_font(r, name=MONO, size=10, color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_run_font(r, bold=True)
        else:
            r = p.add_run(part)
            set_run_font(r)


def add_list_item(doc, text, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r)


def add_heading(doc, level, text):
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    for r in p.runs:
        set_run_font(r, bold=True)


def markdown_to_docx(md_path, docx_path, title):
    doc = Document()
    configure_styles(doc)
    add_footer(doc.sections[0], title)
    add_cover(doc, title, md_path.name)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code))
                code = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(raw)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1 and text == title:
                i += 1
                continue
            if title == "产品总体设计文档" and text.startswith("11. 验收指标"):
                doc.add_page_break()
            add_heading(doc, max(level - 1, 1), text)
            i += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered_match:
            add_list_item(doc, ordered_match.group(1), ordered=True)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            add_list_item(doc, bullet_match.group(1), ordered=False)
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


def main():
    for source, dest, title in SOURCES:
        markdown_to_docx(DELIVERABLES / source, DELIVERABLES / dest, title)
        print(f"created {DELIVERABLES / dest}")


if __name__ == "__main__":
    main()
